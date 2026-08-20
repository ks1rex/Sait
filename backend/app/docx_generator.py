"""
Генерация отчёта в формате ГОСТ Р 7.32 / ГОСТ 2.105.

Структура документа:
  1. Титульный лист
  2. Содержание (TOC-поле; обновить в Word: Ctrl+A → F9)
  3. Введение
  4. Исходные данные
  5. Разделы расчёта (Heading 1/2 по section.level)
  6. Графическая часть (placeholder)
  7. Заключение
  8. Список использованных источников

Формулы выводятся как текст:
  Символ = формула = подставленные_числа = результат  ед.   (N)
Подстановка чисел выполняется через _substitute_values().
"""
from __future__ import annotations

import keyword
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


from .calc_engine import _BUILTIN_NAMES, render_display_template
from .gost_styles import (
    apply_gost_page_setup,
    apply_gost_paragraph_styles,
    apply_page_numbering,
    apply_table_cell_style,
)
from .schemas import CalculationSpec

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Identifiers skipped during formula value substitution
# (same list the calc engine treats as built-ins — keep one source)
_FORMULA_SKIP: frozenset[str] = frozenset(_BUILTIN_NAMES | set(keyword.kwlist))

# Tab-stop positions in twips (A4, left=3 cm, right=1 cm → content=17 cm)
# 1 cm ≈ 567 twips
_CONTENT_TWIPS = 9639   # 17 cm — right edge of content area
_CENTER_TWIPS  = 4820   # 8.5 cm — centre of content area


# ---------------------------------------------------------------------------
# Number / formula helpers
# ---------------------------------------------------------------------------

def _fmt_number(value: float, rounding: int) -> str:
    """ГОСТ-style number: comma decimal separator, trailing zeros stripped."""
    s = f"{value:.{rounding}f}"
    if '.' in s:
        s = s.rstrip('0').rstrip('.')
    return s.replace('.', ',')


def _substitute_values(formula: str, formatted: dict[str, str]) -> str:
    """
    Replace each variable name in *formula* with its pre-formatted string from
    *formatted*.  Python keywords and math-function names are left as-is.
    Supports both Latin and Cyrillic identifiers.

    Each caller pre-builds *formatted* with per-variable precision so that
    input_data values use their natural decimal places (not the current step's
    rounding), and step results use their own step.rounding.
    """
    def _repl(m: re.Match) -> str:
        name = m.group(0)
        if name in _FORMULA_SKIP:
            return name
        return formatted.get(name, name)

    return re.sub(
        r'\b[A-Za-zА-Яа-яЁё_][A-Za-zА-Яа-яЁё_0-9]*\b',
        _repl,
        formula,
    )


def _natural_decimals(value) -> int:
    """
    Number of significant decimal digits of a numeric value (capped at 10).
    Uses fixed-point formatting so small magnitudes that repr() would render
    in scientific notation (e.g. 1e-05) are still measured correctly.
    """
    if isinstance(value, int):
        return 0
    f = float(value)
    if f == int(f):
        return 0
    # %.10f never uses scientific notation; strip trailing zeros to get the
    # natural precision without the artefacts of binary float repr.
    s = f"{f:.10f}".rstrip('0')
    if '.' in s:
        return len(s.split('.')[1])
    return 0


# Одиночная '*' (не часть '**' — возведение в степень) -> ГОСТ-точка умножения.
_MUL_RE = re.compile(r'(?<!\*)\*(?!\*)')


def _display_ops(text: str) -> str:
    """Печатный вид операторов: '*' -> центрированная точка умножения."""
    return _MUL_RE.sub('·', text)


_UNIT_FRAC_RE = re.compile(r'^([^/]+)/([^/]+)$')


def _unit_to_frac(unit: str) -> str:
    """'м³/ч' -> '\\frac{м³}{ч}' — единица измерения внутри m:oMath рисуется
    настоящей дробью, как и остальная формула, а не текстом со слэшем."""
    m = _UNIT_FRAC_RE.match(unit.strip())
    return f'\\frac{{{m.group(1)}}}{{{m.group(2)}}}' if m else unit


# ---------------------------------------------------------------------------
# Native Word equation objects (m:oMath) — what Alt+= builds when you type a
# formula and press space/enter, produced here directly so the user never has
# to do that by hand. python-docx already knows the 'm' (OOXML math) prefix,
# so this needs no new dependency, just hand-built OOXML like the TOC field
# above.
# ---------------------------------------------------------------------------

# Upright (non-italic) in math typesetting convention: digits and operators.
# Everything else (letters — the actual variable/index text) is italic.
# '\t' included so a tab embedded in the formula text (used to push the
# formula number to the right tab stop, still inside the same m:oMath)
# doesn't get wrapped in an italic run.
_UPRIGHT_CHARS = frozenset('0123456789.,=+-·()/ \t')

_BASE_RE = re.compile(r'[A-Za-zА-Яа-яЁё0-9]+$')

# Structural tokens recognized in formula_display / result_symbol text.
# Content inside {...} may itself nest any of these (matched via brace
# depth-counting in _find_matching_brace, not by this regex).
_TOKEN_START_RE = re.compile(r'\\frac\{|\\sqrt\{|_\{|\^\{')


def _auto_subscript(symbol: str) -> str:
    """
    Turn a plain display symbol like 'Qсут' into '_{}'-markup 'Q_{сут}' so it
    renders with a real subscript, same as the rest of a formula_display line
    — matches this project's existing symbol convention (first character is
    the base, the remainder is the index: Qсут, Kmax.ч, Dрасч, Hнас, ...).
    Symbols of length <= 1 (D, h, l, ...) have no index and pass through.
    """
    return symbol if len(symbol) <= 1 else f'{symbol[0]}_{{{symbol[1:]}}}'


def _m(tag: str):
    return OxmlElement(f'm:{tag}')


def _rfonts_cambria_math():
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Cambria Math')
    rFonts.set(qn('w:hAnsi'), 'Cambria Math')
    return rFonts


def _omath_run(text: str, italic: bool):
    r = _m('r')
    rPr = OxmlElement('w:rPr')
    rPr.append(_rfonts_cambria_math())
    if italic:
        rPr.append(OxmlElement('w:i'))
    r.append(rPr)
    t = _m('t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def _m_ctrl_pr(italic: bool = False):
    """m:ctrlPr — the w:rPr Word attaches to structural math elements
    (m:eqArr, m:d, ...), same Cambria Math font as ordinary math runs."""
    ctrlPr = _m('ctrlPr')
    rPr = OxmlElement('w:rPr')
    rPr.append(_rfonts_cambria_math())
    if italic:
        rPr.append(OxmlElement('w:i'))
    ctrlPr.append(rPr)
    return ctrlPr


def _omath_text_runs(text: str) -> list:
    """Split text into m:r runs, alternating italic (letters) / upright
    (digits, operators) — close enough to Word's own math autobuild without
    a full typesetting engine."""
    if not text:
        return []
    runs = []
    start = 0
    cur_italic = text[0] not in _UPRIGHT_CHARS
    for i in range(1, len(text)):
        is_italic = text[i] not in _UPRIGHT_CHARS
        if is_italic != cur_italic:
            runs.append(_omath_run(text[start:i], cur_italic))
            start = i
            cur_italic = is_italic
    runs.append(_omath_run(text[start:], cur_italic))
    return runs


def _find_matching_brace(text: str, open_idx: int) -> int:
    """text[open_idx] must be '{'; returns the index of its matching '}'."""
    depth = 0
    for i in range(open_idx, len(text)):
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
            if depth == 0:
                return i
    raise ValueError(f"formula_display: непарная '{{' ({text!r})")


def _build_omath_children(text: str) -> list:
    """
    Recursively parse *text* into OMML nodes. Recognizes, and lets nest
    inside one another via brace-depth matching:
      \\frac{num}{den}  -> m:f   (a real stacked fraction, not a '/' slash)
      \\sqrt{expr}       -> m:rad (a real radical sign, not literal 'sqrt(...)')
      _{...} / ^{...}   -> m:sSub / m:sSup, base = the identifier run right
                           before the marker (e.g. 'H' in 'H_{нас}') — same
                           as what Word's own Alt+= linear format uses.
    Anything else is plain text, split into italic/upright runs.
    This (plus the sub/superscript structures) is what makes Word render the
    equation in "Professional" (2-D) style — that display is decided purely
    by which OMML elements are present, there is no separate mode flag.
    """
    nodes: list = []
    pending = ''  # plain text since last flush; also the sub/sup base source
    pos = 0
    while pos < len(text):
        m = _TOKEN_START_RE.match(text, pos)
        if not m:
            pending += text[pos]
            pos += 1
            continue

        token = m.group(0)
        brace_open = m.end() - 1  # index of the '{' this token just matched

        if token == '\\frac{':
            num_end = _find_matching_brace(text, brace_open)
            if num_end + 1 >= len(text) or text[num_end + 1] != '{':
                pending += token  # no second group -> not a real \frac, keep literal
                pos = m.end()
                continue
            den_open = num_end + 1
            den_end = _find_matching_brace(text, den_open)

            nodes.extend(_omath_text_runs(pending))
            pending = ''
            f = _m('f')
            num = _m('num')
            num.extend(_build_omath_children(text[brace_open + 1:num_end]))
            f.append(num)
            den = _m('den')
            den.extend(_build_omath_children(text[den_open + 1:den_end]))
            f.append(den)
            nodes.append(f)
            pos = den_end + 1
            continue

        if token == '\\sqrt{':
            end = _find_matching_brace(text, brace_open)

            nodes.extend(_omath_text_runs(pending))
            pending = ''
            rad = _m('rad')
            radPr = _m('radPr')
            deg_hide = _m('degHide')
            deg_hide.set(qn('m:val'), '1')
            radPr.append(deg_hide)
            rad.append(radPr)
            rad.append(_m('deg'))  # empty degree element, required even when hidden
            e = _m('e')
            e.extend(_build_omath_children(text[brace_open + 1:end]))
            rad.append(e)
            nodes.append(rad)
            pos = end + 1
            continue

        # _{...} or ^{...}
        end = _find_matching_brace(text, brace_open)
        base_match = _BASE_RE.search(pending)
        if base_match:
            prefix, base_text = pending[:base_match.start()], base_match.group(0)
        else:
            prefix, base_text = pending, ''

        nodes.extend(_omath_text_runs(prefix))
        pending = ''
        struct = _m('sSub' if token == '_{' else 'sSup')
        e = _m('e')
        e.extend(_omath_text_runs(base_text))
        struct.append(e)
        idx = _m('sub' if token == '_{' else 'sup')
        idx.extend(_build_omath_children(text[brace_open + 1:end]))
        struct.append(idx)
        nodes.append(struct)
        pos = end + 1

    nodes.extend(_omath_text_runs(pending))
    return nodes


def _add_omath(paragraph, text: str) -> None:
    """Append *text* to *paragraph* as a single native Word equation object —
    what you'd get pressing Alt+=, typing the linear form and hitting space,
    done automatically instead of by hand."""
    oMath = _m('oMath')
    for node in _build_omath_children(text):
        oMath.append(node)
    paragraph._p.append(oMath)


def _add_omath_numbered(paragraph, body_text: str, number: int) -> None:
    """
    Append *paragraph* a numbered display equation: same as _add_omath, but
    reproducing the exact structure Word itself builds when you type
    '#(1)' in linear format inside a display equation — m:oMathPara wrapping
    an m:eqArr (maxDist=1) whose single row ends in a '#' run followed by an
    m:d delimiter (auto-parenthesised) around the number. This specific
    construct is what makes Word pin the number to the right margin
    regardless of the paragraph's own alignment — confirmed by inspecting
    the OOXML Word produces after applying "Professional" to a hand-typed
    '#(N)' equation. Word only recognizes '#(N)' while parsing live typed
    linear input, not on already-built OOXML, so this has to be reproduced
    by hand rather than just emitting the literal characters.
    """
    oMathPara = _m('oMathPara')
    oMath = _m('oMath')
    eqArr = _m('eqArr')
    eqArrPr = _m('eqArrPr')
    max_dist = _m('maxDist')
    max_dist.set(qn('m:val'), '1')
    eqArrPr.append(max_dist)
    eqArrPr.append(_m_ctrl_pr(italic=True))
    eqArr.append(eqArrPr)

    e = _m('e')
    e.extend(_build_omath_children(body_text))
    e.append(_omath_run('#', italic=False))

    d = _m('d')
    dPr = _m('dPr')
    dPr.append(_m_ctrl_pr(italic=True))
    d.append(dPr)
    d_e = _m('e')
    d_e.append(_omath_run(str(number), italic=False))
    d.append(d_e)
    e.append(d)

    eqArr.append(e)
    oMath.append(eqArr)
    oMathPara.append(oMath)
    paragraph._p.append(oMathPara)


_PLAIN_SUBSUP_RE = re.compile(r'_\{([^}]*)\}|\^\{([^}]*)\}')


def _add_text_with_subscript(paragraph, text: str) -> None:
    """
    Add *text* (as produced by _auto_subscript, e.g. 'Q_{max.ч}') to a plain
    (non-math) paragraph as normal runs plus real Word subscript/superscript
    runs (font.subscript/superscript). Keeps symbol mentions in running prose
    (e.g. "..., Qmax.ч, рассчитывается по формуле:") in the same register as
    the formula objects, instead of printing them as flat text.
    """
    pos = 0
    for m in _PLAIN_SUBSUP_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        sub, sup = m.group(1), m.group(2)
        run = paragraph.add_run(sub if sub is not None else sup)
        if sub is not None:
            run.font.subscript = True
        else:
            run.font.superscript = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _remove_table_borders(table) -> None:
    """Remove all visible borders from a table via OOXML."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def _add_title_page(doc: Document, spec: CalculationSpec, meta: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get('university', '[Наименование учебного заведения]'))
    r.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run((spec.work_type or 'РАСЧЁТНО-ГРАФИЧЕСКАЯ РАБОТА').upper())
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'по дисциплине «{spec.discipline}»').bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'Тема: {spec.title}')

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"Выполнил: {meta.get('student_name', '[ФИО студента]')}\n")
    p.add_run(f"Группа: {meta.get('group', '[группа]')}\n")
    p.add_run(f"Проверил: {meta.get('supervisor', '[ФИО преподавателя]')}")

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(meta.get('city_year', '[Город, год]'))

    doc.add_page_break()


def _add_toc(doc: Document) -> None:
    """
    Insert auto-collectible TOC field (Heading levels 1-2).
    After opening in Word: press Ctrl+A → F9 to update the field.
    """
    doc.add_heading('Содержание', level=1)

    p = doc.add_paragraph()
    run = p.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)

    hint = doc.add_paragraph(
        '[Откройте в Word и нажмите Ctrl+A → F9 для обновления содержания]'
    )
    hint.paragraph_format.first_line_indent = Cm(0)
    hint.runs[0].italic = True


def _add_intro(doc: Document, spec: CalculationSpec) -> None:
    doc.add_heading('Введение', level=1)
    text = spec.intro_text or '[Введение не задано — заполните вручную]'
    for block in text.split('\n\n'):
        if block.strip():
            doc.add_paragraph(block.strip())


def _add_input_data_table(doc: Document, spec: CalculationSpec) -> None:
    doc.add_heading('Исходные данные', level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0].cells
    headers = ['Обозначение', 'Наименование', 'Значение', 'Ед. изм.']
    for cell, text in zip(hdr, headers):
        cell.text = text
        apply_table_cell_style(cell)
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    for item in spec.input_data:
        row = table.add_row().cells
        # symbol пуст → печатаем id: в документе обозначение обязано совпадать
        # с тем, что стоит в формулах.
        row[0].text = item.symbol or item.id
        row[1].text = item.description
        row[2].text = (
            _fmt_number(float(item.value), 4)
            if isinstance(item.value, (int, float))
            else str(item.value)
        )
        row[3].text = item.unit
        for cell in row:
            apply_table_cell_style(cell)


def _add_formula_row(
    doc: Document,
    step,
    formula_counter: int,
    formatted: dict[str, str],
    symbols: dict[str, str],
) -> None:
    """
    Render the formula line:
        <tab> Symbol = formula = substituted = result unit <tab> (N)
    Uses a centre tab stop at 8.5 cm and a right tab stop at 17 cm,
    so the formula body is centred and the serial number is at the right margin.
    *formatted* maps var_id -> pre-formatted string (each with its own precision).
    *symbols* maps var_id -> display symbol, used only when step.formula_display
    is set (see below).
    """
    value_str = (
        _fmt_number(step.value, step.rounding)
        if step.value is not None
        else '?'
    )
    unit = f' {step.unit}'.rstrip()
    result_symbol = step.result_symbol or step.id

    if step.formula_display:
        # Печатная формула — независимый шаблон, не парсится ни Python, ни
        # asteval, поэтому спецсимволы/индексы в ней не конфликтуют с
        # расчётной formula. {{ id }} подставляется через тот же sandboxed
        # Jinja2, что и intro/conclusion; вся строка целиком становится
        # настоящим объектом-формулой Word (m:oMath) — тем же, что получился
        # бы, если вручную нажать Alt+= и ввести её в линейном виде.
        symbolic = _display_ops(render_display_template(step.formula_display, symbols))
        substituted = _display_ops(render_display_template(step.formula_display, formatted))
        # Ведущее обозначение форматируется тем же авто-индексом, что и
        # остальная формула — иначе оно осталось бы обычным текстом на фоне
        # настоящих подстрочных индексов дальше в строке. Единица измерения —
        # тоже настоящей дробью (м³/ч -> м³ над ч), а не слэшем.
        unit_display = f' {_unit_to_frac(step.unit)}' if step.unit else ''
        full = f'{_auto_subscript(result_symbol)} = {symbolic} = {substituted} = {value_str}{unit_display}'

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(14)
        # m:eqArr + m:d — тот же приём, что Word строит сам для '#(N)' в
        # display-формуле, номер прижимается к правому краю независимо от
        # выравнивания абзаца (см. _add_omath_numbered).
        _add_omath_numbered(p, full, formula_counter)
        return

    p = doc.add_paragraph()
    # Explicit LEFT alignment: formula paragraphs inherit Normal's JUSTIFY,
    # which fights the manual tab stops below and breaks rendering in some
    # converters (e.g. LibreOffice --convert-to pdf drops/garbles the line).
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)

    # Inject tab stops into paragraph XML
    pPr = p._p.get_or_add_pPr()
    tabs_elem = OxmlElement('w:tabs')

    centre_tab = OxmlElement('w:tab')
    centre_tab.set(qn('w:val'), 'center')
    centre_tab.set(qn('w:pos'), str(_CENTER_TWIPS))
    tabs_elem.append(centre_tab)

    right_tab = OxmlElement('w:tab')
    right_tab.set(qn('w:val'), 'right')
    right_tab.set(qn('w:pos'), str(_CONTENT_TWIPS))
    tabs_elem.append(right_tab)

    pPr.append(tabs_elem)

    subst = _display_ops(_substitute_values(step.formula, formatted))
    formula_text = _display_ops(step.formula)
    # step.formula печатается ровно как записана в шаблоне (кроме '*' -> '·') —
    # меняются только цифры в подставленной копии, буквы остаются авторскими.
    formula_line = f'{result_symbol} = {formula_text} = {subst} = {value_str}{unit}'

    # \t → jump to centre tab; second \t → jump to right tab for number
    run = p.add_run(f'\t{formula_line}\t({formula_counter})')
    run.italic = True


def _add_sections(doc: Document, spec: CalculationSpec) -> None:
    formula_counter = 0

    # Pre-build formatted dict: input_data uses natural decimal precision,
    # step results will be added incrementally using their own step.rounding.
    # This prevents input_data values from being rounded to the current step's
    # precision (e.g. T_hl=0.5 showing as "0" when step.rounding=0).
    formatted: dict[str, str] = {}
    for item in spec.input_data:
        try:
            val = float(item.value)
            decimals = _natural_decimals(item.value)
            formatted[item.id] = _fmt_number(val, decimals)
        except (TypeError, ValueError):
            pass

    # symbols: var_id -> display symbol, auto-indexed (_auto_subscript) so
    # {{ id }} placeholders in formula_display render with a real subscript
    # too, not just the parts the template author hand-marks with _{...}.
    # Used only by formula_display rendering.
    symbols: dict[str, str] = {
        item.id: _auto_subscript(item.symbol or item.id) for item in spec.input_data
    }

    for section in spec.sections:
        level = section.level if 1 <= section.level <= 2 else 1
        doc.add_heading(section.title, level=level)

        if section.intro_text:
            p = doc.add_paragraph(section.intro_text)
            p.paragraph_format.first_line_indent = Cm(0)

        for step in section.steps:
            formula_counter += 1

            # Абзац 1: описание величины
            p1 = doc.add_paragraph(f'{step.description}, ')
            _add_text_with_subscript(p1, _auto_subscript(step.result_symbol or step.id))
            p1.add_run(', рассчитывается по формуле:')
            p1.paragraph_format.first_line_indent = Cm(0)

            # Абзац 2: строка формулы с табстопами
            _add_formula_row(doc, step, formula_counter, formatted, symbols)

            # Абзац 3: пояснение «где ...» (если есть)
            if step.explanation:
                p3 = doc.add_paragraph(f'где {step.explanation}')
                p3.paragraph_format.first_line_indent = Cm(0)

            # Add step result to formatted/symbols with its own rounding
            if step.value is not None:
                formatted[step.id] = _fmt_number(step.value, step.rounding)
            symbols[step.id] = _auto_subscript(step.result_symbol or step.id)


def _add_graphics_placeholder(doc: Document) -> None:
    doc.add_heading('Графическая часть', level=1)
    p = doc.add_paragraph(
        '[Графическая часть формируется отдельно — '
        'блок-схема и генплан не входят в автоматический расчёт]'
    )
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].italic = True


def _add_conclusion(doc: Document, spec: CalculationSpec) -> None:
    doc.add_heading('Заключение', level=1)
    doc.add_paragraph(
        spec.conclusion_text or '[Заключение не сгенерировано]'
    )


def _add_references(doc: Document, spec: CalculationSpec) -> None:
    doc.add_heading('Список использованных источников', level=1)
    if spec.references:
        for i, ref in enumerate(spec.references, 1):
            p = doc.add_paragraph(f'{i}. {ref}')
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1.25)
    else:
        p = doc.add_paragraph(
            '[Список источников не определён автоматически — заполните вручную]'
        )
        p.paragraph_format.first_line_indent = Cm(0)
        p.runs[0].italic = True


def _add_chapter_page_breaks(doc: Document) -> None:
    """Insert a page break before every Heading 1 except the first (title-page
    break already precedes it)."""
    seen_first = False
    for p in doc.paragraphs:
        if p.style.name != 'Heading 1':
            continue
        if not seen_first:
            seen_first = True
            continue
        pPr = p._p.get_or_add_pPr()
        pageBreak = OxmlElement('w:pageBreakBefore')
        pageBreak.set(qn('w:val'), '1')
        pPr.append(pageBreak)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_docx(spec: CalculationSpec, meta: dict, output_path: str) -> str:
    """
    Generate a GOST-formatted .docx report.

    Args:
        spec:        Computed CalculationSpec (all step.value must be set).
        meta:        Title-page metadata dict with keys:
                     university, student_name, group, supervisor, city_year.
        output_path: Absolute path to write the .docx file.

    Returns:
        output_path (passthrough for convenience).
    """
    doc = Document()
    apply_gost_page_setup(doc)
    apply_gost_paragraph_styles(doc)
    apply_page_numbering(doc)

    _add_title_page(doc, spec, meta)
    _add_toc(doc)
    _add_intro(doc, spec)
    _add_input_data_table(doc, spec)
    _add_sections(doc, spec)
    _add_graphics_placeholder(doc)
    _add_conclusion(doc, spec)
    _add_references(doc, spec)

    _add_chapter_page_breaks(doc)

    doc.save(output_path)
    return output_path
