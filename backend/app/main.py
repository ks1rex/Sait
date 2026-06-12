"""
FastAPI application — entry point.
All business logic lives in dedicated modules (pdf_extract, calc_engine, etc.).
"""
from __future__ import annotations

import copy
import io
import json
import logging
import os
import subprocess
import tempfile
import uuid
from pathlib import Path
from typing import Annotated, Dict, List, Optional

import fitz
from dotenv import load_dotenv
from docx import Document
from fastapi import Body, Depends, FastAPI, File, Form, HTTPException, Query, UploadFile, status
from fastapi.responses import JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, ValidationError

from .ai_provider import AI_PROVIDER, chat_completion, extract_calculation_spec, extract_variant_inputs, generate_conclusion, minimal_edit_rewrite
from .auth import get_current_user, CurrentUser
from .calc_engine import CalcError, render_text_template, run_calculation
from .docx_md_converter import docx_to_markdown, markdown_to_docx
from .docx_generator import generate_docx
from .gost_styles import apply_gost_styles, remove_toc_section
from .pdf_extract import extract_text_and_tables
from .schemas import CalculationSpec
from .billing import InsufficientTokensError, consume_tokens, get_token_cost
from .supabase_client import get_supabase, get_supabase_as_user
from . import admin as admin_module

load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger("app")

app = FastAPI(
    title="GOST Calculator API",
    description="Backend для генерации расчётных работ по ГОСТ",
    version="0.1.0",
)

_cors_origins = [
    o.strip()
    for o in os.getenv("CORS_ORIGINS", "http://localhost:5173,http://localhost:3000").split(",")
    if o.strip()
]
if "*" in _cors_origins:
    logger.warning(
        "CORS_ORIGINS='*' — все источники разрешены. "
        "В продакшене укажите конкретные домены фронтенда."
    )

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    # The wildcard origin is incompatible with credentials per the CORS spec
    allow_credentials="*" not in _cors_origins,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(admin_module.router)

from fastapi import Request as _Request
from fastapi.responses import JSONResponse as _JSONResponse

@app.exception_handler(InsufficientTokensError)
async def insufficient_tokens_handler(_req: _Request, exc: InsufficientTokensError):
    return _JSONResponse(
        status_code=402,
        content={
            "error": "insufficient_tokens",
            "required": exc.required,
            "balance": exc.balance,
        },
    )

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_PDF_CONTENT_TYPES = {"application/pdf", "application/octet-stream", "binary/octet-stream"}
_TEXT_CONTENT_TYPES = {"text/plain", "text/csv"}

_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
_MANIFEST_PATH = _TEMPLATES_DIR / "manifest.json"
_VALID_GENERATION_MODES = {"universal", "fixed_template", "custom_template"}

_MAX_FILE_BYTES = 20 * 1024 * 1024  # 20 MB


async def _read_limited(upload: UploadFile, label: str) -> bytes:
    data = await upload.read(_MAX_FILE_BYTES + 1)
    if len(data) > _MAX_FILE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail={"error": f"{label}: файл превышает 20 МБ"},
        )
    return data



# ---------------------------------------------------------------------------
# Request / Response models
# ---------------------------------------------------------------------------

class HealthResponse(BaseModel):
    status: str
    version: str


class UploadedFile(BaseModel):
    file_type: str       # task | methodology | variant_data
    storage_path: str
    original_name: str


class UploadResponse(BaseModel):
    project_id: str
    files: List[UploadedFile]
    task_text_length: int


class ExtractResponse(BaseModel):
    project_id: str
    spec: CalculationSpec


class ComputeRequest(BaseModel):
    project_id: str


class ComputeResponse(BaseModel):
    project_id: str
    results: Dict[str, float]


class GenerateMeta(BaseModel):
    university: str = ""
    student_name: str = ""
    group: str = ""
    supervisor: str = ""
    city_year: str = ""


class GenerateRequest(BaseModel):
    project_id: str
    meta: GenerateMeta = Field(default_factory=GenerateMeta)


class GenerateResponse(BaseModel):
    project_id: str
    docx_url: str
    pdf_url: Optional[str] = None
    warning: Optional[str] = None


class TemplateInfo(BaseModel):
    id: str
    title: str
    discipline: str = ""
    work_type: str = ""
    description: str = ""
    spec_file: str


class ChatMessageIn(BaseModel):
    message: str


class ChatMessageOut(BaseModel):
    id: str
    role: str
    content: str
    created_at: str


class ChatResponse(BaseModel):
    reply: str
    docx_url: str
    pdf_url: Optional[str] = None


class MeResponse(BaseModel):
    token_balance: int
    unlimited_access: bool
    is_admin: bool = False


class RedeemCodeRequest(BaseModel):
    code: str


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _require_project(project_id: str, user_id: str) -> dict:
    """Loads a project row and verifies ownership. Raises 404 if not found."""
    db = get_supabase()
    result = (
        db.table("projects")
        .select("*")
        .eq("id", project_id)
        .eq("user_id", user_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": "Проект не найден"},
        )
    return result.data[0]


def _validate_pdf_bytes(pdf_bytes: bytes, label: str) -> None:
    """Raises HTTP 400 if bytes are not a readable PDF with at least one page."""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page_count = doc.page_count
        doc.close()
    except fitz.FileDataError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": f"{label}: невалидный или повреждённый PDF-файл"},
        )
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": f"{label}: не удалось открыть файл как PDF"},
        )
    if page_count == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": f"{label}: PDF не содержит страниц"},
        )


def _bytes_to_text(file_bytes: bytes, filename: str, content_type: str | None) -> str:
    """Convert uploaded file bytes to plain text (handles PDF and TXT)."""
    is_text = (
        content_type in _TEXT_CONTENT_TYPES
        or (filename or "").lower().endswith(".txt")
    )
    if is_text:
        return file_bytes.decode("utf-8", errors="replace")

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        return extract_text_and_tables(tmp_path)
    finally:
        os.unlink(tmp_path)


def _download_and_extract(storage_path: str) -> str:
    """Download a file from 'uploads' bucket and extract its text."""
    db = get_supabase()
    try:
        file_bytes: bytes = db.storage.from_("uploads").download(storage_path)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Не удалось скачать файл {storage_path}: {exc}"},
        )
    ext = storage_path.rsplit(".", 1)[-1].lower() if "." in storage_path else ""
    content_type = "text/plain" if ext == "txt" else "application/pdf"
    return _bytes_to_text(file_bytes, storage_path, content_type)


def _validate_template_id(template_id: str) -> None:
    """Check that template_id exists in manifest.json and its spec file is on disk."""
    if not _MANIFEST_PATH.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": "Список шаблонов (manifest.json) не найден на сервере"},
        )
    try:
        manifest = json.loads(_MANIFEST_PATH.read_text(encoding="utf-8"))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Не удалось прочитать список шаблонов: {exc}"},
        )
    entry = next((item for item in manifest if item.get("id") == template_id), None)
    if entry is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": f"Шаблон с id='{template_id}' не найден в manifest.json"},
        )
    spec_file = entry.get("spec_file", "")
    if not (_TEMPLATES_DIR / "specs" / spec_file).exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": f"Файл спецификации для шаблона '{template_id}' не найден на сервере"},
        )


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.get("/health", response_model=HealthResponse, tags=["system"])
async def health() -> HealthResponse:
    return HealthResponse(status="ok", version=app.version)


@app.get("/me", response_model=MeResponse, tags=["billing"])
async def get_me(user: CurrentUser) -> MeResponse:
    """Возвращает баланс токенов и флаг безлимитного доступа текущего пользователя."""
    db = get_supabase()
    try:
        result = (
            db.table("profiles")
            .select("token_balance, unlimited_access, is_admin")
            .eq("id", user["user_id"])
            .single()
            .execute()
        )
        data = result.data
    except Exception:
        # .single() raises when the profile row is missing — report 404, not 500
        data = None
    if not data:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND,
                            detail={"error": "Профиль не найден"})
    return MeResponse(**data)


@app.post("/redeem-code", tags=["billing"])
async def redeem_code(body: RedeemCodeRequest, user: CurrentUser):
    """Активирует код доступа, зачисляет токены на баланс."""
    user_client = get_supabase_as_user(user["jwt"])
    try:
        result = user_client.rpc("redeem_code", {"p_code": body.code}).execute()
    except Exception as exc:
        msg = str(exc)
        if "invalid_or_used_code" in msg:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": "Код не найден или уже использован"},
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка активации кода: {exc}"},
        )
    new_balance: int = result.data if isinstance(result.data, int) else 0
    return {"token_balance": new_balance}


@app.get("/templates", response_model=List[TemplateInfo], tags=["templates"])
async def list_templates(user: CurrentUser) -> List[TemplateInfo]:
    """
    Возвращает список доступных фиксированных шаблонов расчёта из
    templates/manifest.json. Используется в режиме fixed_template.
    """
    if not _MANIFEST_PATH.exists():
        return []
    try:
        data = json.loads(_MANIFEST_PATH.read_text(encoding="utf-8"))
        # Only expose templates whose spec_file is present on disk
        return [
            TemplateInfo(**item)
            for item in data
            if (_TEMPLATES_DIR / "specs" / item.get("spec_file", "")).exists()
        ]
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Не удалось прочитать список шаблонов: {exc}"},
        )


_VALID_SUB_MODES = {"format_only", "minimal_edit", "chat"}


@app.post("/upload", response_model=UploadResponse, tags=["projects"])
async def upload_pdf(
    task: Annotated[
        Optional[UploadFile],
        File(description="PDF: задание (для universal/fixed_template)"),
    ] = None,
    methodology: Annotated[
        Optional[UploadFile], File(description="PDF: методичка (опционально)")
    ] = None,
    variant_data: Annotated[
        Optional[UploadFile], File(description="PDF или TXT: исходные данные по варианту (опционально)")
    ] = None,
    template: Annotated[
        Optional[UploadFile],
        File(description="docx или PDF: пользовательский шаблон (только для custom_template)"),
    ] = None,
    generation_mode: Annotated[
        str,
        Form(description="Режим генерации: universal | fixed_template | custom_template"),
    ] = "universal",
    template_id: Annotated[
        Optional[str],
        Form(description="ID шаблона из GET /templates (только для fixed_template)"),
    ] = None,
    sub_mode: Annotated[
        Optional[str],
        Form(description="Под-режим для custom_template: format_only | minimal_edit | chat"),
    ] = None,
    user: CurrentUser = None,
) -> UploadResponse:
    """
    Принимает до трёх файлов: задание (обязательно), методичка и
    исходные данные по варианту (оба опциональны).
    Сохраняет каждый в Storage, создаёт записи в project_files.

    generation_mode=fixed_template + template_id указывает, что для этого
    проекта будет использоваться готовый зашитый шаблон расчёта.
    """
    if generation_mode not in _VALID_GENERATION_MODES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "error": (
                    f"Недопустимый generation_mode '{generation_mode}'. "
                    f"Допустимые значения: {sorted(_VALID_GENERATION_MODES)}"
                )
            },
        )

    if generation_mode == "fixed_template" and not template_id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": "Для режима fixed_template необходимо указать template_id"},
        )

    if generation_mode == "fixed_template":
        _validate_template_id(template_id)  # type: ignore[arg-type]
        if methodology and methodology.filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": (
                        "В режиме fixed_template принимается только один файл (task). "
                        "Поле 'methodology' должно быть пустым."
                    )
                },
            )
        if variant_data and variant_data.filename:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": (
                        "В режиме fixed_template принимается только один файл (task). "
                        "Поле 'variant_data' должно быть пустым."
                    )
                },
            )

    # ── custom_template ──────────────────────────────────────────────────────
    if generation_mode == "custom_template":
        effective_sub_mode = sub_mode or "format_only"
        if effective_sub_mode not in _VALID_SUB_MODES:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": f"Недопустимый sub_mode '{effective_sub_mode}'. Допустимые: {sorted(_VALID_SUB_MODES)}"},
            )
        # chat allows both files to be absent (generation from scratch)
        # minimal_edit: template required, task required
        # format_only:  template required, task ignored
        if effective_sub_mode in ("format_only", "minimal_edit") and (not template or not template.filename):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": "Для custom_template необходимо передать файл template (.docx или .pdf)"},
            )
        if effective_sub_mode == "minimal_edit" and (not task or not task.filename):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": "Для sub_mode='minimal_edit' необходим файл task (PDF или TXT с новым заданием)"},
            )

        # ── Process template file (present for format_only/minimal_edit; optional for chat) ──
        tpl_filename: str | None = None
        ct_docx_bytes: bytes | None = None
        if template and template.filename:
            tpl_filename = template.filename
            tpl_ext = os.path.splitext(tpl_filename)[1].lower()
            tpl_is_pdf = tpl_ext == ".pdf" or template.content_type in _PDF_CONTENT_TYPES
            tpl_is_docx = tpl_ext == ".docx" or template.content_type == _DOCX_CONTENT_TYPE
            if not (tpl_is_pdf or tpl_is_docx):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail={"error": "Шаблон должен быть файлом .docx или .pdf"},
                )
            tpl_bytes = await _read_limited(template, "Шаблон")
            if not tpl_bytes:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail={"error": "Файл шаблона пустой"},
                )
            if tpl_is_pdf:
                _validate_pdf_bytes(tpl_bytes, "Шаблон")
                try:
                    with tempfile.TemporaryDirectory() as tmp_conv:
                        pdf_tmp = os.path.join(tmp_conv, "input.pdf")
                        with open(pdf_tmp, "wb") as f:
                            f.write(tpl_bytes)
                        subprocess.run(
                            [_SOFFICE_BIN, "--headless", "--convert-to", "docx", "--outdir", tmp_conv, pdf_tmp],
                            check=True, capture_output=True, timeout=_SOFFICE_TIMEOUT,
                        )
                        docx_tmp = os.path.join(tmp_conv, "input.docx")
                        if not os.path.exists(docx_tmp):
                            raise RuntimeError("LibreOffice не создал .docx")
                        with open(docx_tmp, "rb") as f:
                            ct_docx_bytes = f.read()
                except FileNotFoundError:
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail={"error": "LibreOffice (soffice) не установлен — не удалось конвертировать PDF в docx"},
                    )
                except (subprocess.CalledProcessError, subprocess.TimeoutExpired, RuntimeError) as exc:
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail={"error": f"Ошибка конвертации PDF → docx: {exc}"},
                    )
            else:
                ct_docx_bytes = tpl_bytes

        # ── Process task file (required for minimal_edit; optional for chat) ──
        ct_task_bytes: bytes | None = None
        ct_task_filename: str | None = None
        if task and task.filename and effective_sub_mode in ("minimal_edit", "chat"):
            ct_task_bytes = await _read_limited(task, "Задание")
            if not ct_task_bytes:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail={"error": "Файл задания (task) пустой"},
                )
            ct_task_filename = task.filename

        # ── Upload to Storage ──────────────────────────────────────────────────
        ct_project_id = str(uuid.uuid4())
        ct_user_id: str = user["user_id"]
        ct_db = get_supabase()

        ct_storage_path: str | None = None
        if ct_docx_bytes:
            ct_storage_path = f"{ct_user_id}/{ct_project_id}/template.docx"
            try:
                ct_db.storage.from_("uploads").upload(
                    path=ct_storage_path,
                    file=ct_docx_bytes,
                    file_options={"content-type": _DOCX_CONTENT_TYPE},
                )
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail={"error": f"Ошибка загрузки шаблона в хранилище: {exc}"},
                )

        ct_task_storage_path: str | None = None
        if ct_task_bytes and ct_task_filename:
            task_ext = os.path.splitext(ct_task_filename)[1].lower()
            task_is_txt = task_ext == ".txt" or (task and task.content_type in _TEXT_CONTENT_TYPES)
            ct_task_storage_path = f"{ct_user_id}/{ct_project_id}/task{'.txt' if task_is_txt else '.pdf'}"
            task_ct = "text/plain" if task_is_txt else "application/pdf"
            try:
                ct_db.storage.from_("uploads").upload(
                    path=ct_task_storage_path,
                    file=ct_task_bytes,
                    file_options={"content-type": task_ct},
                )
            except Exception as exc:
                if ct_storage_path:
                    try:
                        ct_db.storage.from_("uploads").remove([ct_storage_path])
                    except Exception:
                        pass
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail={"error": f"Ошибка загрузки задания в хранилище: {exc}"},
                )

        ct_title = os.path.splitext(tpl_filename)[0] if tpl_filename else "Новый документ"
        try:
            ct_db.table("projects").insert({
                "id": ct_project_id,
                "user_id": ct_user_id,
                "title": ct_title,
                "status": "uploaded",
                "generation_mode": "custom_template",
            }).execute()
        except Exception as exc:
            paths_to_remove = [p for p in [ct_storage_path, ct_task_storage_path] if p]
            if paths_to_remove:
                try:
                    ct_db.storage.from_("uploads").remove(paths_to_remove)
                except Exception:
                    pass
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": f"Ошибка создания проекта в БД: {exc}"},
            )

        ct_record: dict = {"project_id": ct_project_id, "sub_mode": effective_sub_mode}
        if ct_storage_path:
            ct_record["source_storage_path"] = ct_storage_path
        if ct_task_storage_path:
            ct_record["task_storage_path"] = ct_task_storage_path
        try:
            ct_db.table("custom_templates").insert(ct_record).execute()
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": f"Ошибка записи custom_templates: {exc}"},
            )

        uploaded_files = []
        if ct_storage_path and tpl_filename:
            uploaded_files.append(UploadedFile(
                file_type="template",
                storage_path=ct_storage_path,
                original_name=tpl_filename,
            ))
        if ct_task_storage_path and ct_task_filename:
            uploaded_files.append(UploadedFile(
                file_type="task",
                storage_path=ct_task_storage_path,
                original_name=ct_task_filename,
            ))

        return UploadResponse(
            project_id=ct_project_id,
            files=uploaded_files,
            task_text_length=0,
        )

    # ── universal / fixed_template ────────────────────────────────────────────
    if not task or not task.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": f"Для режима '{generation_mode}' необходимо передать файл task (PDF)"},
        )

    if task.content_type not in _PDF_CONTENT_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": "Поле 'task' должно содержать PDF-файл"},
        )

    task_bytes = await _read_limited(task, "Задание")
    if not task_bytes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": "Файл задания пустой"},
        )

    _validate_pdf_bytes(task_bytes, "Задание")

    task_text = _bytes_to_text(task_bytes, task.filename or "", task.content_type)
    if not task_text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "error": (
                    "Не удалось извлечь текст из PDF задания. "
                    "Возможно, файл отсканирован без распознавания (OCR)."
                )
            },
        )

    # Read optional files
    methodology_bytes: bytes | None = None
    if methodology and methodology.filename:
        methodology_bytes = await _read_limited(methodology, "Методичка")
        if methodology_bytes:
            _validate_pdf_bytes(methodology_bytes, "Методичка")
        else:
            methodology_bytes = None

    variant_data_bytes: bytes | None = None
    if variant_data and variant_data.filename:
        variant_data_bytes = await _read_limited(variant_data, "Исходные данные")
        if not variant_data_bytes:
            variant_data_bytes = None

    # Generate project ID and upload files to Storage
    project_id = str(uuid.uuid4())
    user_id: str = user["user_id"]
    db = get_supabase()

    uploaded: list[UploadedFile] = []
    storage_paths: list[str] = []  # for rollback

    def _upload_one(file_bytes: bytes, file_type: str, suffix: str, content_type: str, original_name: str) -> UploadedFile:
        path = f"{user_id}/{project_id}/{file_type}{suffix}"
        try:
            db.storage.from_("uploads").upload(
                path=path,
                file=file_bytes,
                file_options={"content-type": content_type},
            )
        except Exception as exc:
            if storage_paths:
                try:
                    db.storage.from_("uploads").remove(storage_paths)
                except Exception:
                    pass
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": f"Ошибка загрузки файла {file_type} в хранилище: {exc}"},
            )
        storage_paths.append(path)
        return UploadedFile(file_type=file_type, storage_path=path, original_name=original_name)

    uploaded.append(_upload_one(task_bytes, "task", ".pdf", "application/pdf", task.filename or "task.pdf"))

    if methodology_bytes:
        uploaded.append(_upload_one(methodology_bytes, "methodology", ".pdf", "application/pdf", methodology.filename or "methodology.pdf"))  # type: ignore[union-attr]

    if variant_data_bytes:
        vd_filename = variant_data.filename or "variant_data"  # type: ignore[union-attr]
        is_txt = (variant_data.content_type in _TEXT_CONTENT_TYPES or vd_filename.lower().endswith(".txt"))  # type: ignore[union-attr]
        vd_suffix = ".txt" if is_txt else ".pdf"
        vd_ct = "text/plain" if is_txt else "application/pdf"
        uploaded.append(_upload_one(variant_data_bytes, "variant_data", vd_suffix, vd_ct, vd_filename))

    # Create project record
    title = os.path.splitext(task.filename or "Без названия")[0]
    project_row: dict = {
        "id": project_id,
        "user_id": user_id,
        "title": title,
        "status": "uploaded",
        "generation_mode": generation_mode,
    }
    if template_id:
        project_row["template_id"] = template_id
    try:
        db.table("projects").insert(project_row).execute()
    except Exception as exc:
        try:
            db.storage.from_("uploads").remove(storage_paths)
        except Exception:
            pass
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка создания проекта в БД: {exc}"},
        )

    # Create project_files records
    try:
        db.table("project_files").insert(
            [
                {
                    "project_id": project_id,
                    "file_type": f.file_type,
                    "storage_path": f.storage_path,
                    "original_name": f.original_name,
                }
                for f in uploaded
            ]
        ).execute()
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка записи метаданных файлов: {exc}"},
        )

    return UploadResponse(
        project_id=project_id,
        files=uploaded,
        task_text_length=len(task_text),
    )


@app.post("/extract", response_model=ExtractResponse, tags=["projects"])
async def extract_spec(
    project_id: Annotated[str, Query(description="ID проекта, полученный из /upload")],
    user: CurrentUser,
) -> ExtractResponse:
    """
    Строит CalculationSpec для проекта в зависимости от generation_mode:
    - universal: AI извлекает структуру из загруженных PDF (текущая логика).
    - fixed_template / custom_template: заглушка, будет реализована в Блоке 3-6.
    """
    db = get_supabase()
    user_id: str = user["user_id"]

    # 1. Verify project ownership and load project row (contains generation_mode)
    project = _require_project(project_id, user_id)
    generation_mode: str = project.get("generation_mode", "universal")

    # 1b. Determine sub_mode for billing (custom_template only)
    _extract_sub_mode: str | None = None
    if generation_mode == "custom_template":
        _ct_meta = (
            db.table("custom_templates")
            .select("sub_mode")
            .eq("project_id", project_id)
            .execute()
        )
        _extract_sub_mode = (_ct_meta.data[0]["sub_mode"] if _ct_meta.data else None)

    # 1c. Consume tokens BEFORE any AI/processing work
    _cost = get_token_cost(generation_mode, _extract_sub_mode)
    consume_tokens(
        get_supabase_as_user(user["jwt"]),
        amount=_cost,
        reason=f"extract:{generation_mode}" + (f":{_extract_sub_mode}" if _extract_sub_mode else ""),
        project_id=project_id,
    )

    # ── fixed_template ───────────────────────────────────────────────────────
    if generation_mode == "fixed_template":
        template_id: str | None = project.get("template_id")
        if not template_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": "Проект не имеет template_id для режима fixed_template"},
            )

        spec_path = _TEMPLATES_DIR / "specs" / f"{template_id}.json"
        if not spec_path.exists():
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": f"Файл спецификации для шаблона '{template_id}' не найден"},
            )
        spec_dict = copy.deepcopy(json.loads(spec_path.read_text(encoding="utf-8")))

        files_result = (
            db.table("project_files")
            .select("file_type, storage_path")
            .eq("project_id", project_id)
            .execute()
        )
        if not files_result.data:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": "К проекту не привязано ни одного файла"},
            )
        files_by_type: dict[str, str] = {
            row["file_type"]: row["storage_path"] for row in files_result.data
        }
        if "task" not in files_by_type:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": "Файл варианта (task) не найден в проекте"},
            )

        variant_text = _download_and_extract(files_by_type["task"])
        if not variant_text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={"error": "Не удалось извлечь текст из PDF варианта"},
            )

        input_schema = [
            {
                "id": item["id"],
                "symbol": item.get("symbol", ""),
                "description": item.get("description", ""),
                "unit": item.get("unit", ""),
            }
            for item in spec_dict.get("input_data", [])
        ]

        try:
            variant_result = extract_variant_inputs(input_schema, variant_text)
        except ValueError as exc:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail={"error": f"AI вернул некорректный JSON: {exc}"},
            )
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail={"error": f"Ошибка при обращении к AI: {exc}"},
            )

        for item in spec_dict["input_data"]:
            if item["id"] in variant_result.overrides:
                item["value"] = variant_result.overrides[item["id"]]

        try:
            spec = CalculationSpec.model_validate(spec_dict)
        except ValidationError as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail={
                    "error": "Шаблон не прошёл валидацию после применения варианта",
                    "validation_errors": exc.errors(),
                },
            )

        try:
            db.table("calculation_specs").upsert(
                {"project_id": project_id, "spec_json": spec.model_dump()},
                on_conflict="project_id",
            ).execute()
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": f"Ошибка сохранения спецификации: {exc}"},
            )

        db.table("projects").update({"status": "extracted"}).eq("id", project_id).execute()

        try:
            db.table("ai_usage").insert(
                {
                    "user_id": user_id,
                    "project_id": project_id,
                    "provider": AI_PROVIDER,
                    "model": variant_result.model,
                    "input_tokens": variant_result.input_tokens,
                    "output_tokens": variant_result.output_tokens,
                }
            ).execute()
        except Exception:
            pass

        return ExtractResponse(project_id=project_id, spec=spec)

    # ── custom_template ───────────────────────────────────────────────────────
    if generation_mode == "custom_template":
        ct_row = (
            db.table("custom_templates")
            .select("source_storage_path, task_storage_path, sub_mode")
            .eq("project_id", project_id)
            .execute()
        )
        if not ct_row.data:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error": "Запись custom_templates не найдена для этого проекта"},
            )
        ct_info = ct_row.data[0]
        ct_sub_mode: str = ct_info.get("sub_mode", "format_only")
        ct_source_path: str = ct_info["source_storage_path"]

        # ── helpers shared by both sub_modes ─────────────────────────────────

        def _ct_save_and_upload(doc_obj: Document) -> tuple[bytes, bytes | None]:
            """Save doc to temp dir, convert to PDF. Return (docx_bytes, pdf_bytes|None)."""
            with tempfile.TemporaryDirectory() as _tmp:
                _out_docx = os.path.join(_tmp, "report.docx")
                doc_obj.save(_out_docx)
                with open(_out_docx, "rb") as _f:
                    _docx_b = _f.read()
                _pdf_b: bytes | None = None
                try:
                    subprocess.run(
                        [_SOFFICE_BIN, "--headless", "--convert-to", "pdf", "--outdir", _tmp, _out_docx],
                        check=True, capture_output=True, timeout=_SOFFICE_TIMEOUT,
                    )
                    _out_pdf = os.path.join(_tmp, "report.pdf")
                    if os.path.exists(_out_pdf):
                        with open(_out_pdf, "rb") as _f:
                            _pdf_b = _f.read()
                except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
                    pass
            return _docx_b, _pdf_b

        def _ct_upload_outputs(docx_b: bytes, pdf_b: bytes | None) -> tuple[str, str | None]:
            """Upload to outputs bucket, return (docx_signed_url, pdf_signed_url|None)."""
            _docx_path = f"{user_id}/{project_id}/report.docx"
            _pdf_path = f"{user_id}/{project_id}/report.pdf"
            try:
                db.storage.from_("outputs").upload(
                    path=_docx_path, file=docx_b,
                    file_options={"content-type": _DOCX_CONTENT_TYPE},
                )
            except Exception as _exc:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail={"error": f"Ошибка загрузки docx в outputs: {_exc}"},
                )
            if pdf_b:
                try:
                    db.storage.from_("outputs").upload(
                        path=_pdf_path, file=pdf_b,
                        file_options={"content-type": "application/pdf"},
                    )
                except Exception:
                    pass
            try:
                _docx_url = db.storage.from_("outputs").create_signed_url(_docx_path, 3600)["signedURL"]
            except Exception:
                _docx_url = ""
            _pdf_url: str | None = None
            if pdf_b:
                try:
                    _pdf_url = db.storage.from_("outputs").create_signed_url(_pdf_path, 3600)["signedURL"]
                except Exception:
                    pass
            _upd: dict = {"status": "done", "output_docx_path": _docx_path}
            if pdf_b:
                _upd["output_pdf_path"] = _pdf_path
            db.table("projects").update(_upd).eq("id", project_id).execute()
            return _docx_url, _pdf_url

        # ── format_only ───────────────────────────────────────────────────────
        if ct_sub_mode == "format_only":
            try:
                ct_docx_bytes: bytes = db.storage.from_("uploads").download(ct_source_path)
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail={"error": f"Не удалось скачать шаблон из хранилища: {exc}"},
                )
            try:
                doc = Document(io.BytesIO(ct_docx_bytes))
                apply_gost_styles(doc)
                remove_toc_section(doc)
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail={"error": f"Ошибка применения ГОСТ-стилей: {exc}"},
                )
            fo_docx_b, fo_pdf_b = _ct_save_and_upload(doc)
            fo_docx_url, fo_pdf_url = _ct_upload_outputs(fo_docx_b, fo_pdf_b)
            return JSONResponse(content={
                "project_id": project_id,
                "mode": "format_only",
                "message": "Документ отформатирован по ГОСТ, доступен для скачивания",
                "docx_url": fo_docx_url,
                "pdf_url": fo_pdf_url,
            })

        # ── minimal_edit ──────────────────────────────────────────────────────
        if ct_sub_mode == "minimal_edit":
            ct_task_path: str | None = ct_info.get("task_storage_path")
            if not ct_task_path:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail={"error": "Путь к файлу задания (task_storage_path) не найден в custom_templates"},
                )

            # Download both files from Storage
            try:
                me_tpl_bytes: bytes = db.storage.from_("uploads").download(ct_source_path)
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail={"error": f"Не удалось скачать шаблон: {exc}"},
                )
            try:
                me_task_bytes: bytes = db.storage.from_("uploads").download(ct_task_path)
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail={"error": f"Не удалось скачать файл задания: {exc}"},
                )

            # Convert template docx → Markdown
            try:
                tpl_doc = Document(io.BytesIO(me_tpl_bytes))
                template_md = docx_to_markdown(tpl_doc)
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail={"error": f"Ошибка конвертации шаблона в Markdown: {exc}"},
                )

            # Extract text from task file
            task_ext = ct_task_path.rsplit(".", 1)[-1].lower() if "." in ct_task_path else ""
            task_ct = "text/plain" if task_ext == "txt" else "application/pdf"
            try:
                me_task_text = _bytes_to_text(me_task_bytes, ct_task_path, task_ct)
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail={"error": f"Ошибка извлечения текста из задания: {exc}"},
                )
            if not me_task_text.strip():
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail={"error": "Не удалось извлечь текст из файла задания (OCR?)"},
                )

            # Call DeepSeek (complex rewrite task — uses fallback/pro model)
            try:
                me_result = minimal_edit_rewrite(template_md, me_task_text)
            except ValueError as exc:
                raise HTTPException(
                    status_code=status.HTTP_502_BAD_GATEWAY,
                    detail={"error": f"AI вернул некорректный ответ: {exc}"},
                )
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_502_BAD_GATEWAY,
                    detail={"error": f"Ошибка при обращении к AI: {exc}"},
                )

            # Build new docx from AI markdown response
            try:
                new_doc = Document()
                apply_gost_styles(new_doc)
                markdown_to_docx(me_result.markdown, new_doc)
            except Exception as exc:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail={"error": f"Ошибка сборки docx из Markdown: {exc}"},
                )

            me_docx_b, me_pdf_b = _ct_save_and_upload(new_doc)
            me_docx_url, me_pdf_url = _ct_upload_outputs(me_docx_b, me_pdf_b)

            # Log AI token usage (non-fatal)
            try:
                db.table("ai_usage").insert({
                    "user_id": user_id,
                    "project_id": project_id,
                    "provider": AI_PROVIDER,
                    "model": me_result.model,
                    "input_tokens": me_result.input_tokens,
                    "output_tokens": me_result.output_tokens,
                }).execute()
            except Exception:
                pass

            return JSONResponse(content={
                "project_id": project_id,
                "mode": "minimal_edit",
                "message": "Документ переработан с минимальными изменениями, доступен для скачивания",
                "docx_url": me_docx_url,
                "pdf_url": me_pdf_url,
            })

        if ct_sub_mode == "chat":
            return JSONResponse(content={
                "project_id": project_id,
                "mode": "chat",
                "message": "Используйте POST /chat/{project_id} для интерактивного редактирования",
            })

        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail={"error": f"sub_mode='{ct_sub_mode}' ещё не реализован"},
        )

    # 2. Load file records for this project
    files_result = (
        db.table("project_files")
        .select("file_type, storage_path")
        .eq("project_id", project_id)
        .execute()
    )
    if not files_result.data:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": "К проекту не привязано ни одного файла"},
        )

    files_by_type: dict[str, str] = {
        row["file_type"]: row["storage_path"] for row in files_result.data
    }

    if "task" not in files_by_type:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": "Файл задания (task) не найден в проекте"},
        )

    # 3. Download and extract text from each file
    task_text = _download_and_extract(files_by_type["task"])
    if not task_text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": "Не удалось извлечь текст из файла задания"},
        )

    methodology_text: str | None = None
    if "methodology" in files_by_type:
        methodology_text = _download_and_extract(files_by_type["methodology"]) or None

    extra_inputs_text: str | None = None
    if "variant_data" in files_by_type:
        extra_inputs_text = _download_and_extract(files_by_type["variant_data"]) or None

    # 4. Call AI (primary model)
    try:
        ai_result = extract_calculation_spec(
            task_text=task_text,
            methodology_text=methodology_text,
            extra_inputs_text=extra_inputs_text,
            use_fallback_model=False,
        )
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail={"error": f"AI вернул некорректный JSON: {exc}"},
        )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail={"error": f"Ошибка при обращении к AI: {exc}"},
        )

    total_input_tokens = ai_result.input_tokens
    total_output_tokens = ai_result.output_tokens
    model_used = ai_result.model

    # 5. Validate with Pydantic; retry on fallback model if validation fails
    try:
        spec = CalculationSpec.model_validate(ai_result.spec_dict)
    except ValidationError:
        try:
            fallback = extract_calculation_spec(
                task_text=task_text,
                methodology_text=methodology_text,
                extra_inputs_text=extra_inputs_text,
                use_fallback_model=True,
            )
            total_input_tokens += fallback.input_tokens
            total_output_tokens += fallback.output_tokens
            model_used = fallback.model
            spec = CalculationSpec.model_validate(fallback.spec_dict)
        except ValidationError as exc:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail={
                    "error": (
                        "Модель вернула некорректную структуру спецификации. "
                        "Попробуйте ещё раз или загрузите другой PDF."
                    ),
                    "validation_errors": exc.errors(),
                },
            )
        except ValueError as exc:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail={"error": f"Fallback-модель вернула некорректный JSON: {exc}"},
            )
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail={"error": f"Ошибка при обращении к fallback-модели: {exc}"},
            )

    # 6. Upsert spec into calculation_specs
    try:
        db.table("calculation_specs").upsert(
            {"project_id": project_id, "spec_json": spec.model_dump()},
            on_conflict="project_id",
        ).execute()
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка сохранения спецификации: {exc}"},
        )

    # 7. Update project status
    db.table("projects").update({"status": "extracted"}).eq("id", project_id).execute()

    # 8. Log token usage (non-fatal)
    try:
        db.table("ai_usage").insert(
            {
                "user_id": user_id,
                "project_id": project_id,
                "provider": AI_PROVIDER,
                "model": model_used,
                "input_tokens": total_input_tokens,
                "output_tokens": total_output_tokens,
            }
        ).execute()
    except Exception:
        pass

    return ExtractResponse(project_id=project_id, spec=spec)


class ProjectMeta(BaseModel):
    id: str
    title: str
    status: str
    generation_mode: str
    created_at: str
    output_docx_url: Optional[str] = None
    output_pdf_url: Optional[str] = None


@app.get("/project/{project_id}", response_model=ProjectMeta, tags=["projects"])
async def get_project(project_id: str, user: CurrentUser) -> ProjectMeta:
    """Возвращает метаданные проекта и свежие подписанные URL для скачивания файлов."""
    project = _require_project(project_id, user["user_id"])
    db = get_supabase()
    _TTL = 3600

    def _signed(path: Optional[str]) -> Optional[str]:
        if not path:
            return None
        try:
            r = db.storage.from_("outputs").create_signed_url(path, _TTL)
            return r.get("signedURL")
        except Exception:
            return None

    return ProjectMeta(
        id=project["id"],
        title=project.get("title", ""),
        status=project.get("status", "pending"),
        generation_mode=project.get("generation_mode", "universal"),
        created_at=project.get("created_at", ""),
        output_docx_url=_signed(project.get("output_docx_path")),
        output_pdf_url=_signed(project.get("output_pdf_path")),
    )


@app.get("/spec/{project_id}", tags=["projects"])
async def get_spec(project_id: str, user: CurrentUser):
    """
    Возвращает сохранённую спецификацию расчёта для проекта.
    Для custom_template/format_only возвращает JSON с mode и message вместо CalculationSpec.
    """
    project = _require_project(project_id, user["user_id"])
    generation_mode: str = project.get("generation_mode", "universal")

    db = get_supabase()

    if generation_mode == "custom_template":
        ct_row = (
            db.table("custom_templates")
            .select("sub_mode")
            .eq("project_id", project_id)
            .execute()
        )
        sub_mode = ct_row.data[0]["sub_mode"] if ct_row.data else "format_only"
        _ct_messages = {
            "format_only": "Документ отформатирован по ГОСТ, доступен для скачивания",
            "minimal_edit": "Документ переработан с минимальными изменениями, доступен для скачивания",
            "chat": "Интерактивный режим — используйте POST /chat/{project_id} для редактирования",
        }
        if sub_mode in _ct_messages:
            return JSONResponse(content={"mode": sub_mode, "message": _ct_messages[sub_mode]})

    result = (
        db.table("calculation_specs")
        .select("spec_json")
        .eq("project_id", project_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": "Спецификация ещё не создана. Сначала выполните /extract."},
        )
    return CalculationSpec.model_validate(result.data[0]["spec_json"])


@app.put("/spec/{project_id}", response_model=CalculationSpec, tags=["projects"])
async def update_spec(project_id: str, spec: CalculationSpec, user: CurrentUser) -> CalculationSpec:
    """
    Сохраняет отредактированную пользователем спецификацию.
    Экран проверки/редактирования — обязательный шаг перед /compute.
    """
    _require_project(project_id, user["user_id"])

    db = get_supabase()
    db.table("calculation_specs").update(
        {"spec_json": spec.model_dump()}
    ).eq("project_id", project_id).execute()

    return spec


# ---------------------------------------------------------------------------
# Chat constants & helpers
# ---------------------------------------------------------------------------

_CHAT_SYSTEM_PROMPT = (
    "Ты помогаешь редактировать техническую работу. ВСЕГДА оформляй "
    "результат по ГОСТ (Times New Roman 14pt, поля 30/10/20/20мм, "
    "межстрочный интервал 1.5, выравнивание по ширине, заголовки с "
    "отступом 1.25см) — ЕСЛИ пользователь явно не попросил иначе.\n"
    "Формулы нумеруй сквозно в скобках по порядку появления: "
    "`Q = V / t  (1)`. "
    "Нумерацию формул сохраняй и обновляй при каждом редактировании.\n"
    "Все остальные требования по содержанию, структуре, формулировкам — "
    "бери из сообщений пользователя.\n"
    "Отвечай строго в следующем формате:\n"
    "1. Markdown-представление ПОЛНОГО актуального текста документа "
    "(не diff, а целиком). Используй # для Заголовка 1, ## для Заголовка 2, "
    "| для таблиц.\n"
    "2. Строка-разделитель: <!-- REPLY -->\n"
    "3. Короткий комментарий (1-2 предложения): что именно изменено или добавлено."
)

# Фразы, при наличии которых в истории apply_gost_styles НЕ применяется
_NO_GOST_PHRASES = (
    "не по гост", "без гост", "другой шрифт", "другое форматирование",
    "другой стиль", "arial", "calibri", "verdana", "без отступ",
)


def _should_apply_gost(chat_history: list[dict]) -> bool:
    """Return False if any user message contains an explicit non-GOST formatting request."""
    for msg in chat_history:
        if msg.get("role") == "user":
            text = msg.get("content", "").lower()
            if any(phrase in text for phrase in _NO_GOST_PHRASES):
                return False
    return True


def _parse_chat_response(full_content: str) -> tuple[str, str]:
    """
    Split AI response on <!-- REPLY --> marker.
    Returns (doc_markdown, reply_comment).
    If marker absent, whole content is used as doc_markdown.
    """
    parts = full_content.split("<!-- REPLY -->", 1)
    doc_md = parts[0].strip()
    reply = parts[1].strip() if len(parts) > 1 else "Документ обновлён."
    return doc_md, reply


# ---------------------------------------------------------------------------
# Chat endpoints
# ---------------------------------------------------------------------------

@app.post("/chat/{project_id}", response_model=ChatResponse, tags=["chat"])
async def chat_turn(
    project_id: str,
    body: ChatMessageIn,
    user: CurrentUser,
) -> ChatResponse:
    """
    One turn of the interactive document-editing chat.
    Saves user message, calls AI with full context, updates the draft docx/pdf,
    saves assistant reply, returns signed URLs to the new draft.
    """
    db = get_supabase()
    user_id: str = user["user_id"]

    # 1. Verify project ownership and generation_mode
    project = _require_project(project_id, user_id)
    if project.get("generation_mode") != "custom_template":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": "Чат доступен только для проектов с generation_mode='custom_template'"},
        )

    # 2. Load custom_templates record
    ct_row = (
        db.table("custom_templates")
        .select("source_storage_path, task_storage_path, sub_mode")
        .eq("project_id", project_id)
        .execute()
    )
    if not ct_row.data:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": "Запись custom_templates не найдена"},
        )
    ct_info = ct_row.data[0]
    ct_source_path: str | None = ct_info.get("source_storage_path")
    ct_task_path: str | None = ct_info.get("task_storage_path")

    # 3. Load existing chat history (ordered by created_at)
    hist_result = (
        db.table("chat_messages")
        .select("id, role, content, created_at")
        .eq("project_id", project_id)
        .order("created_at")
        .execute()
    )
    history: list[dict] = hist_result.data or []
    is_first_turn = not history

    # 4. Consume tokens BEFORE saving message (insufficient → 402, nothing written)
    consume_tokens(
        get_supabase_as_user(user["jwt"]),
        amount=get_token_cost("custom_template", "chat"),
        reason="chat_turn",
        project_id=project_id,
    )

    # 5. Save user message to DB BEFORE calling AI (so retry is safe)
    try:
        db.table("chat_messages").insert({
            "project_id": project_id,
            "role": "user",
            "content": body.message,
        }).execute()
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка сохранения сообщения: {exc}"},
        )

    # 5. Build messages list for AI
    ai_messages: list[dict] = [{"role": "system", "content": _CHAT_SYSTEM_PROMPT}]

    # On first turn (or always): inject template + task as context
    # These are NOT stored in DB — they're re-injected each turn for consistency
    if ct_source_path:
        try:
            tpl_bytes = db.storage.from_("uploads").download(ct_source_path)
            tpl_doc = Document(io.BytesIO(tpl_bytes))
            template_md = docx_to_markdown(tpl_doc)
            ai_messages.append({
                "role": "user",
                "content": f"Исходный шаблон документа (Markdown):\n\n{template_md}",
            })
            ai_messages.append({
                "role": "assistant",
                "content": "Исходный шаблон принят. Готов работать с документом.",
            })
        except Exception:
            pass  # template unavailable — continue without it

    if ct_task_path:
        try:
            task_bytes = db.storage.from_("uploads").download(ct_task_path)
            task_ext = ct_task_path.rsplit(".", 1)[-1].lower() if "." in ct_task_path else ""
            task_ct = "text/plain" if task_ext == "txt" else "application/pdf"
            task_text = _bytes_to_text(task_bytes, ct_task_path, task_ct)
            if task_text.strip():
                ai_messages.append({
                    "role": "user",
                    "content": f"Условие / задание:\n\n{task_text}",
                })
                ai_messages.append({
                    "role": "assistant",
                    "content": "Задание принято.",
                })
        except Exception:
            pass

    # Append real chat history (user/assistant turns from DB)
    for msg in history:
        ai_messages.append({"role": msg["role"], "content": msg["content"]})

    # Current user message
    ai_messages.append({"role": "user", "content": body.message})

    # 6. Call AI
    try:
        ai_result = chat_completion(ai_messages)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail={"error": f"Ошибка при обращении к AI: {exc}"},
        )

    full_response = ai_result.content
    doc_markdown, reply_comment = _parse_chat_response(full_response)

    # 7. Save assistant message to DB
    try:
        db.table("chat_messages").insert({
            "project_id": project_id,
            "role": "assistant",
            "content": full_response,
        }).execute()
    except Exception:
        pass  # non-fatal — docx still gets built

    # 8. Build docx from AI markdown response
    apply_gost = _should_apply_gost(history + [{"role": "user", "content": body.message}])
    try:
        chat_doc = Document()
        if apply_gost:
            apply_gost_styles(chat_doc)
        markdown_to_docx(doc_markdown, chat_doc)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка сборки docx из ответа AI: {exc}"},
        )

    # 9. Save docx + pdf to outputs (upsert — overwrite previous draft)
    draft_docx_path = f"{user_id}/{project_id}/chat_draft.docx"
    draft_pdf_path = f"{user_id}/{project_id}/chat_draft.pdf"

    with tempfile.TemporaryDirectory() as tmp_chat:
        chat_docx_file = os.path.join(tmp_chat, "chat_draft.docx")
        chat_doc.save(chat_docx_file)
        with open(chat_docx_file, "rb") as f:
            chat_docx_bytes = f.read()

        chat_pdf_bytes: bytes | None = None
        try:
            subprocess.run(
                [_SOFFICE_BIN, "--headless", "--convert-to", "pdf", "--outdir", tmp_chat, chat_docx_file],
                check=True, capture_output=True, timeout=_SOFFICE_TIMEOUT,
            )
            chat_pdf_file = os.path.join(tmp_chat, "chat_draft.pdf")
            if os.path.exists(chat_pdf_file):
                with open(chat_pdf_file, "rb") as f:
                    chat_pdf_bytes = f.read()
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
            pass

    try:
        db.storage.from_("outputs").upload(
            path=draft_docx_path,
            file=chat_docx_bytes,
            file_options={"content-type": _DOCX_CONTENT_TYPE, "upsert": "true"},
        )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка загрузки черновика в хранилище: {exc}"},
        )

    if chat_pdf_bytes:
        try:
            db.storage.from_("outputs").upload(
                path=draft_pdf_path,
                file=chat_pdf_bytes,
                file_options={"content-type": "application/pdf", "upsert": "true"},
            )
        except Exception:
            pass

    # 10. Generate signed URLs and update project
    try:
        docx_url = db.storage.from_("outputs").create_signed_url(draft_docx_path, 3600)["signedURL"]
    except Exception:
        docx_url = ""

    pdf_url: str | None = None
    if chat_pdf_bytes:
        try:
            pdf_url = db.storage.from_("outputs").create_signed_url(draft_pdf_path, 3600)["signedURL"]
        except Exception:
            pass

    upd: dict = {"status": "computed", "output_docx_path": draft_docx_path}
    if chat_pdf_bytes:
        upd["output_pdf_path"] = draft_pdf_path
    db.table("projects").update(upd).eq("id", project_id).execute()

    # 11. Log token usage (non-fatal)
    try:
        db.table("ai_usage").insert({
            "user_id": user_id,
            "project_id": project_id,
            "provider": AI_PROVIDER,
            "model": ai_result.model,
            "input_tokens": ai_result.input_tokens,
            "output_tokens": ai_result.output_tokens,
        }).execute()
    except Exception:
        pass

    return ChatResponse(reply=reply_comment, docx_url=docx_url, pdf_url=pdf_url)


@app.get("/chat/{project_id}", response_model=List[ChatMessageOut], tags=["chat"])
async def get_chat_history(project_id: str, user: CurrentUser) -> List[ChatMessageOut]:
    """Возвращает всю историю сообщений чата для проекта в хронологическом порядке."""
    _require_project(project_id, user["user_id"])
    db = get_supabase()
    result = (
        db.table("chat_messages")
        .select("id, role, content, created_at")
        .eq("project_id", project_id)
        .order("created_at")
        .execute()
    )
    return [
        ChatMessageOut(
            id=row["id"],
            role=row["role"],
            content=row["content"],
            created_at=row["created_at"],
        )
        for row in (result.data or [])
    ]


@app.post("/compute", response_model=ComputeResponse, tags=["projects"])
async def compute(
    project_id: Annotated[str, Query(description="ID проекта")],
    user: CurrentUser,
) -> ComputeResponse:
    """
    Запускает расчётный движок по текущей (возможно, отредактированной)
    спецификации. Сохраняет step.value и conclusion_text обратно в БД.
    Не обращается к AI для расчёта — только для генерации заключения.
    """
    db = get_supabase()
    user_id: str = user["user_id"]

    # 1. Verify ownership
    _require_project(project_id, user_id)

    # 2. Load current spec (includes any user edits from PUT /spec)
    result = (
        db.table("calculation_specs")
        .select("spec_json")
        .eq("project_id", project_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": "Спецификация не найдена. Сначала выполните /extract."},
        )
    spec = CalculationSpec.model_validate(result.data[0]["spec_json"])

    # 3. Run calculation engine (mutates spec.sections[*].steps[*].value in place)
    try:
        results = run_calculation(spec)
    except CalcError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={"error": str(exc), "step_id": exc.step_id},
        )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка расчёта: {exc}"},
        )

    # 4. Render Jinja2 templates (fixed_template) or generate via AI (universal)
    if spec.intro_text_template:
        try:
            spec.intro_text = render_text_template(spec.intro_text_template, spec, results)
        except Exception:
            pass  # non-fatal; intro_text stays None

    if spec.conclusion_text_template:
        try:
            spec.conclusion_text = render_text_template(spec.conclusion_text_template, spec, results)
        except Exception:
            spec.conclusion_text = None
    else:
        try:
            spec.conclusion_text = generate_conclusion(spec.model_dump(), results)
        except Exception:
            spec.conclusion_text = None

    # 5. Persist updated spec (with computed values and conclusion)
    db.table("calculation_specs").update(
        {"spec_json": spec.model_dump()}
    ).eq("project_id", project_id).execute()

    # 6. Update project status
    db.table("projects").update({"status": "computed"}).eq("id", project_id).execute()

    return ComputeResponse(project_id=project_id, results=results)


_SOFFICE_TIMEOUT = 90  # seconds
# Resolves soffice executable: LIBREOFFICE_PATH env var → system PATH fallback.
# Windows example: C:\Program Files\LibreOffice\program\soffice.exe
# Linux/Docker:    soffice  (installed via apt, already on PATH)
_SOFFICE_BIN = os.getenv("LIBREOFFICE_PATH", "soffice")


@app.post("/generate", response_model=GenerateResponse, tags=["projects"])
async def generate(
    project_id: Annotated[str, Query(description="ID проекта")],
    user: CurrentUser,
    meta: Annotated[GenerateMeta, Body()] = GenerateMeta(),
) -> GenerateResponse:
    """
    Генерирует .docx через docx_generator и конвертирует в .pdf через
    LibreOffice headless. Загружает оба файла в bucket «outputs».
    Если LibreOffice недоступен — возвращает только docx с предупреждением
    в поле warning.
    """
    db = get_supabase()
    user_id: str = user["user_id"]

    # 1. Verify ownership and check status
    project = _require_project(project_id, user_id)
    if project["status"] not in ("computed", "done"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "error": (
                    "Расчёт ещё не выполнен. "
                    "Сначала выполните /compute."
                )
            },
        )

    # 2. Load computed spec
    result = (
        db.table("calculation_specs")
        .select("spec_json")
        .eq("project_id", project_id)
        .execute()
    )
    if not result.data:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail={"error": "Спецификация не найдена."},
        )
    spec = CalculationSpec.model_validate(result.data[0]["spec_json"])

    # 3. Guard: all steps must have computed values
    uncomputed = [
        step.id
        for section in spec.sections
        for step in section.steps
        if step.value is None
    ]
    if uncomputed:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={
                "error": "Не все шаги вычислены. Выполните /compute заново.",
                "uncomputed_steps": uncomputed,
            },
        )

    storage_prefix = f"{user_id}/{project_id}"
    docx_storage_path = f"{storage_prefix}/report.docx"
    pdf_storage_path: Optional[str] = None
    pdf_warning: Optional[str] = None

    with tempfile.TemporaryDirectory() as tmp_dir:
        docx_path = os.path.join(tmp_dir, "report.docx")

        # 4. Generate .docx
        try:
            generate_docx(spec, meta.model_dump(), docx_path)
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": f"Ошибка генерации docx: {exc}"},
            )

        # 5. Convert to .pdf via LibreOffice headless (optional)
        pdf_path = os.path.join(tmp_dir, "report.pdf")
        try:
            subprocess.run(
                [
                    _SOFFICE_BIN,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmp_dir,
                    docx_path,
                ],
                check=True,
                capture_output=True,
                timeout=_SOFFICE_TIMEOUT,
            )
            if not os.path.exists(pdf_path):
                pdf_warning = "LibreOffice не создал PDF-файл (проверьте установку)."
                pdf_path = None
        except FileNotFoundError:
            pdf_warning = (
                "LibreOffice (soffice) не установлен на сервере. "
                "PDF-версия недоступна — скачайте .docx и конвертируйте самостоятельно."
            )
            pdf_path = None
        except subprocess.TimeoutExpired:
            pdf_warning = f"Конвертация в PDF превысила таймаут ({_SOFFICE_TIMEOUT} с)."
            pdf_path = None
        except subprocess.CalledProcessError as exc:
            stderr = exc.stderr.decode(errors="replace")[:300] if exc.stderr else ""
            pdf_warning = f"Ошибка LibreOffice при конвертации: {stderr}"
            pdf_path = None

        # 6. Upload .docx to Storage (upsert — allow regeneration)
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
        try:
            db.storage.from_("outputs").upload(
                path=docx_storage_path,
                file=docx_bytes,
                file_options={
                    "content-type": (
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document"
                    ),
                    "upsert": "true",
                },
            )
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={"error": f"Ошибка загрузки docx в хранилище: {exc}"},
            )

        # 7. Upload .pdf to Storage (if conversion succeeded)
        if pdf_path and os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            try:
                pdf_storage_path = f"{storage_prefix}/report.pdf"
                db.storage.from_("outputs").upload(
                    path=pdf_storage_path,
                    file=pdf_bytes,
                    file_options={"content-type": "application/pdf", "upsert": "true"},
                )
            except Exception as exc:
                pdf_warning = (pdf_warning or "") + f" Ошибка загрузки PDF: {exc}"
                pdf_storage_path = None

    # 8. Create signed URLs (1 hour)
    _SIGNED_URL_TTL = 3600
    try:
        docx_signed = db.storage.from_("outputs").create_signed_url(
            docx_storage_path, _SIGNED_URL_TTL
        )
        docx_url: str = docx_signed.get("signedURL", "")
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Не удалось создать ссылку для скачивания docx: {exc}"},
        )

    pdf_url: Optional[str] = None
    if pdf_storage_path:
        try:
            pdf_signed = db.storage.from_("outputs").create_signed_url(
                pdf_storage_path, _SIGNED_URL_TTL
            )
            pdf_url = pdf_signed.get("signedURL")
        except Exception:
            pass  # signed URL failure is non-fatal for PDF

    # 9. Persist output paths and mark project done
    db.table("projects").update(
        {
            "output_docx_path": docx_storage_path,
            "output_pdf_path": pdf_storage_path,
            "status": "done",
        }
    ).eq("id", project_id).execute()

    return GenerateResponse(
        project_id=project_id,
        docx_url=docx_url,
        pdf_url=pdf_url,
        warning=pdf_warning,
    )


_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_DOCX_EXTENSIONS = {".docx"}


@app.post("/format-gost", tags=["tools"])
async def format_gost(
    file: Annotated[UploadFile, File(description=".docx файл для приведения к ГОСТ")],
    include_toc: Annotated[
        bool,
        Query(description="Сохранить раздел 'Содержание' (true) или удалить (false)"),
    ] = True,
    user: CurrentUser = None,  # type: ignore[assignment]  # FastAPI Depends overrides default
) -> Response:
    """
    Применяет ГОСТ-стили (поля, шрифты, межстрочный интервал, стили ячеек таблиц)
    к загруженному .docx файлу.

    Если include_toc=false — удаляет заголовок 'Содержание' и поле TOC из документа.
    Если include_toc=true и поля TOC нет — стили применяются без его добавления.

    Возвращает обработанный .docx файл для скачивания.
    Не требует project_id — работает как standalone утилита для любого документа.
    """
    filename = file.filename or "document.docx"
    ext = os.path.splitext(filename)[1].lower()

    is_docx = (
        ext in _DOCX_EXTENSIONS
        or file.content_type == _DOCX_CONTENT_TYPE
    )
    if not is_docx:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": "Поддерживаются только файлы .docx"},
        )

    # Consume tokens BEFORE processing — if insufficient, return 402 immediately
    if user:
        consume_tokens(
            get_supabase_as_user(user["jwt"]),
            amount=get_token_cost("format_gost"),
            reason="format_gost",
        )

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": "Загруженный файл пустой"},
        )

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error": f"Не удалось открыть файл как .docx: {exc}"},
        )

    try:
        apply_gost_styles(doc)
        if not include_toc:
            remove_toc_section(doc)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error": f"Ошибка применения стилей: {exc}"},
        )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = os.path.splitext(filename)[0] + "_gost.docx"
    return Response(
        content=buf.read(),
        media_type=_DOCX_CONTENT_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )
