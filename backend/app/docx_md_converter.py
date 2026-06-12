"""
Bidirectional converter: python-docx Document <-> Markdown text.

docx_to_markdown(doc)        – extracts text preserving headings and tables
markdown_to_docx(md, doc)   – adds parsed markdown content to an existing doc
"""
from __future__ import annotations

import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .gost_styles import apply_table_cell_style


# ---------------------------------------------------------------------------
# docx → Markdown
# ---------------------------------------------------------------------------

def _iter_block_items(doc: Document):
    """Yield ('p', Paragraph) or ('t', Table) in document body order."""
    body = doc.element.body
    for child in body:
        if child.tag == qn('w:p'):
            yield 'p', Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield 't', Table(child, doc)


def _heading_level(para: Paragraph) -> int | None:
    """Return 1 or 2 if para has a heading style, else None."""
    name = (para.style.name or '') if para.style else ''
    # English: "Heading 1", "Heading 2"
    if name.startswith('Heading '):
        try:
            return min(int(name.split()[-1]), 2)
        except ValueError:
            return None
    # Russian: "Заголовок 1", "Заголовок 2"
    if 'аголовок' in name:
        parts = name.split()
        for part in parts:
            try:
                return min(int(part), 2)
            except ValueError:
                continue
    # Fall back to outline level in paragraph XML
    pPr = para._p.find(qn('w:pPr'))
    if pPr is not None:
        lvl_el = pPr.find(qn('w:outlineLvl'))
        if lvl_el is not None:
            val = lvl_el.get(qn('w:val'))
            if val is not None:
                return min(int(val) + 1, 2)
    return None


def _table_to_markdown(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ').replace('|', r'\|') for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ''
    n_cols = max(len(r) for r in rows)
    normalized = [r + [''] * (n_cols - len(r)) for r in rows]
    lines = ['| ' + ' | '.join(normalized[0]) + ' |',
             '| ' + ' | '.join(['---'] * n_cols) + ' |']
    for row in normalized[1:]:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines)


def docx_to_markdown(doc: Document) -> str:
    """
    Convert a python-docx Document to Markdown text.
    Heading 1/2 → # / ##, tables → GFM pipe tables, other paragraphs → plain text.
    Empty paragraphs and page-break-only paragraphs are skipped.
    """
    parts: list[str] = []
    for kind, block in _iter_block_items(doc):
        if kind == 'p':
            para: Paragraph = block
            text = para.text.strip()
            if not text:
                continue
            # Skip paragraphs that are purely page-break markers
            has_pb = any(
                br.get(qn('w:type')) == 'page'
                for br in para._p.findall('.//' + qn('w:br'))
            )
            if has_pb and not text:
                continue
            lvl = _heading_level(para)
            if lvl == 1:
                parts.append(f'# {text}')
            elif lvl == 2:
                parts.append(f'## {text}')
            else:
                parts.append(text)
        else:
            table: Table = block
            md_table = _table_to_markdown(table)
            if md_table:
                parts.append(md_table)
    return '\n\n'.join(parts)


# ---------------------------------------------------------------------------
# Markdown → docx
# ---------------------------------------------------------------------------

_TABLE_ROW_RE = re.compile(r'^\s*\|')
_TABLE_SEP_RE = re.compile(r'^\s*\|[\s\-:|]+\|\s*$')


def _add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    """Parse GFM pipe-table lines and add a Table to doc."""
    data_rows = [line for line in table_lines if not _TABLE_SEP_RE.match(line)]
    rows = []
    for line in data_rows:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(n_cols):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = row_data[c_idx] if c_idx < len(row_data) else ''
            apply_table_cell_style(cell)
    # Make first row bold (header)
    if rows:
        for cell in tbl.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True


def markdown_to_docx(md_text: str, doc: Document) -> None:
    """
    Parse *md_text* and append content to *doc*.
    Mapping: # → Heading 1, ## / ### → Heading 2,
             GFM table → docx Table (Table Grid style),
             everything else → Normal paragraph.
    Assumes heading/paragraph styles are already configured on *doc*.
    """
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Table block: collect all consecutive | lines
        if _TABLE_ROW_RE.match(line):
            table_lines: list[str] = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        stripped = line.strip()

        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped:
            doc.add_paragraph(stripped)
        # empty line → skip

        i += 1
