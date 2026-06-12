"""
Fast smoke test for key endpoints — no AI calls (cheap, run after every change).
Covers: /health, /me, /templates, /redeem-code (negative), /upload (universal),
GET /spec (404 before extract), /format-gost, /admin/* (403 + 200).

Run: cd backend && venv/Scripts/python smoke_test.py
"""
from __future__ import annotations

import io
import os
import sys
import time
import uuid

import fitz
import httpx
from docx import Document
from dotenv import load_dotenv
from jose import jwt as jose_jwt

load_dotenv()

BASE = "http://localhost:8000"
SUPABASE_URL = os.environ["SUPABASE_URL"]
SUPABASE_KEY = os.environ["SUPABASE_SERVICE_ROLE_KEY"]
JWT_SECRET = os.environ["SUPABASE_JWT_SECRET"]

CLIENT = httpx.Client(trust_env=False, timeout=60)
ADMIN_HEADERS = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"}

_failures: list[str] = []


def check(label: str, ok: bool, info: str = "") -> None:
    print(f"{'OK  ' if ok else 'FAIL'} [{label}] {info}")
    if not ok:
        _failures.append(label)


def ensure_user(email: str) -> str:
    r = CLIENT.post(
        f"{SUPABASE_URL}/auth/v1/admin/users",
        headers=ADMIN_HEADERS,
        json={"email": email, "password": "smoke-test-pw-9x!", "email_confirm": True},
    )
    if r.status_code in (200, 201):
        return r.json()["id"]
    r2 = CLIENT.get(f"{SUPABASE_URL}/auth/v1/admin/users?per_page=200", headers=ADMIN_HEADERS)
    for u in r2.json().get("users", []):
        if u.get("email") == email:
            return u["id"]
    print(f"cannot create/find user {email}: {r.status_code} {r.text[:200]}")
    sys.exit(1)


def set_profile_flag(user_id: str, **flags) -> None:
    r = CLIENT.patch(
        f"{SUPABASE_URL}/rest/v1/profiles?id=eq.{user_id}",
        headers={**ADMIN_HEADERS, "Content-Type": "application/json", "Prefer": "return=minimal"},
        json=flags,
    )
    assert r.status_code in (200, 204), f"profile patch failed: {r.status_code} {r.text[:200]}"


def make_token(user_id: str, email: str) -> str:
    now = int(time.time())
    return jose_jwt.encode(
        {"sub": user_id, "email": email, "role": "authenticated",
         "aud": "authenticated", "exp": now + 3600, "iat": now},
        JWT_SECRET, algorithm="HS256",
    )


def tiny_pdf_bytes() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Smoke test PDF. Q = 100 m3/day. Variant 1.")
    data = doc.tobytes()
    doc.close()
    return data


def tiny_docx_bytes() -> bytes:
    d = Document()
    d.add_heading("Smoke heading", level=1)
    d.add_paragraph("Smoke paragraph for format-gost.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def main() -> None:
    # users
    uid = ensure_user("smoke-user@local.dev")
    set_profile_flag(uid, unlimited_access=True)
    tok = make_token(uid, "smoke-user@local.dev")
    H = {"Authorization": f"Bearer {tok}"}

    admin_uid = ensure_user("smoke-admin@local.dev")
    set_profile_flag(admin_uid, is_admin=True)
    admin_tok = make_token(admin_uid, "smoke-admin@local.dev")
    HA = {"Authorization": f"Bearer {admin_tok}"}

    # 1. health
    r = CLIENT.get(f"{BASE}/health")
    check("GET /health", r.status_code == 200 and r.json().get("status") == "ok")

    # 2. unauthenticated → 401/403
    r = CLIENT.get(f"{BASE}/me")
    check("GET /me no-auth", r.status_code in (401, 403), f"status={r.status_code}")

    # 3. /me
    r = CLIENT.get(f"{BASE}/me", headers=H)
    check("GET /me", r.status_code == 200 and r.json().get("unlimited_access") is True,
          f"{r.status_code} {r.text[:120]}")

    # 4. /templates
    r = CLIENT.get(f"{BASE}/templates", headers=H)
    check("GET /templates", r.status_code == 200 and isinstance(r.json(), list))

    # 5. bogus redeem code → 404
    r = CLIENT.post(f"{BASE}/redeem-code", headers=H, json={"code": "NO-SUCH-CODE"})
    check("POST /redeem-code bogus", r.status_code == 404, f"status={r.status_code}")

    # 6. upload universal
    r = CLIENT.post(
        f"{BASE}/upload", headers=H,
        data={"generation_mode": "universal"},
        files={"task": ("task.pdf", tiny_pdf_bytes(), "application/pdf")},
    )
    check("POST /upload universal", r.status_code == 200, f"{r.status_code} {r.text[:200]}")
    project_id = r.json().get("project_id") if r.status_code == 200 else None

    # 7. spec before extract → 404
    if project_id:
        r = CLIENT.get(f"{BASE}/spec/{project_id}", headers=H)
        check("GET /spec before extract", r.status_code == 404, f"status={r.status_code}")

        # 8. foreign project → 404
        r = CLIENT.get(f"{BASE}/spec/{project_id}", headers=HA)
        check("GET /spec foreign project", r.status_code == 404, f"status={r.status_code}")

        # 9. compute without spec → 404
        r = CLIENT.post(f"{BASE}/compute", headers=H, params={"project_id": project_id})
        check("POST /compute no spec", r.status_code == 404, f"status={r.status_code}")

    # 10. invalid project id
    r = CLIENT.get(f"{BASE}/project/{uuid.uuid4()}", headers=H)
    check("GET /project bogus id", r.status_code == 404, f"status={r.status_code}")

    # 11. upload invalid mode
    r = CLIENT.post(
        f"{BASE}/upload", headers=H,
        data={"generation_mode": "bogus"},
        files={"task": ("task.pdf", tiny_pdf_bytes(), "application/pdf")},
    )
    check("POST /upload bad mode", r.status_code == 400, f"status={r.status_code}")

    # 12. upload non-PDF as task
    r = CLIENT.post(
        f"{BASE}/upload", headers=H,
        data={"generation_mode": "universal"},
        files={"task": ("task.pdf", b"not a pdf at all", "application/pdf")},
    )
    check("POST /upload corrupt pdf", r.status_code == 400, f"status={r.status_code}")

    # 13. format-gost
    r = CLIENT.post(
        f"{BASE}/format-gost", headers=H,
        files={"file": ("doc.docx", tiny_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    check("POST /format-gost", r.status_code == 200 and len(r.content) > 1000,
          f"{r.status_code} {len(r.content)} bytes")

    # 14. format-gost wrong type
    r = CLIENT.post(f"{BASE}/format-gost", headers=H,
                    files={"file": ("doc.txt", b"hello", "text/plain")})
    check("POST /format-gost wrong type", r.status_code == 400, f"status={r.status_code}")

    # 15. admin endpoints: non-admin → 403
    r = CLIENT.get(f"{BASE}/admin/stats", headers=H)
    check("GET /admin/stats non-admin", r.status_code == 403, f"status={r.status_code}")

    # 16. admin endpoints: admin → 200
    r = CLIENT.get(f"{BASE}/admin/stats", headers=HA)
    check("GET /admin/stats admin", r.status_code == 200, f"{r.status_code} {r.text[:120]}")

    r = CLIENT.get(f"{BASE}/admin/users", headers=HA)
    check("GET /admin/users admin", r.status_code == 200)

    print()
    if _failures:
        print(f"SMOKE FAILED: {len(_failures)} checks: {_failures}")
        sys.exit(1)
    print("SMOKE OK: all checks passed")


if __name__ == "__main__":
    main()
